# -*- coding: utf-8 -*-
"""生成毕业设计报告 Word 文档"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
WORKSPACE_ROOT = os.path.dirname(PROJECT_ROOT)
OUTPUT = os.path.join(WORKSPACE_ROOT, 'word', 'report_draft.docx')

doc = Document()

# ==================== 样式设置 ====================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Pt(24)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.first_line_indent = None
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(15)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(6)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

# ==================== 辅助函数 ====================
def add_para(text, bold=False, font_size=None, align=None, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if align:
        p.alignment = align
    if not indent:
        p.paragraph_format.first_line_indent = Pt(0)
    return p

def add_code(text):
    """添加代码块（等宽字体，灰色背景效果用缩进模拟）"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return p

def add_image_placeholder(caption):
    """添加图片占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(f'【此处插入图片：{caption}】')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Pt(0)
    r = cap.add_run(f'图：{caption}')
    r.font.size = Pt(10.5)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].first_line_indent = Pt(0)
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10.5)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            cell.paragraphs[0].first_line_indent = Pt(0)
    return table

# ==================== 封面 ====================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据库原理与应用')
run.font.size = Pt(22)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('课 程 设 计 报 告 书')
run.font.size = Pt(26)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('课题名称：豆瓣电影数据分析系统')
run.font.size = Pt(16)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

for label in ['姓    名：', '学    号：', '院    系：', '专业班级：', '指导教师：']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(label + '________________')
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ==================== 目录（占位） ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('目    录')
run.font.size = Pt(18)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
p.paragraph_format.first_line_indent = Pt(0)

doc.add_paragraph()
add_para('（生成 .docx 后，在 Word 中插入 → 引用 → 目录，自动生成）', indent=False)
doc.add_page_break()

# ==================== 引言 ====================
doc.add_heading('引  言', level=1)
add_para('随着互联网的飞速发展，电影产业与在线平台积累了海量的数据资源。豆瓣电影作为中国最具影响力的电影评分社区，汇聚了数千万用户的评分、评论和观影记录。如何高效地存储、管理并分析这些数据，从中挖掘出有价值的观影规律和行业趋势，是数据库应用领域的重要课题。')
add_para('传统的电影信息管理方式主要依赖人工检索和简单排序，难以满足多维度数据分析与可视化展示的需求。随着数据库技术和Web可视化技术的成熟，构建一套集数据采集、存储、分析和可视化展示于一体的电影数据分析系统成为可能。本系统利用MySQL数据库对豆瓣电影数据进行规范化存储，通过Python爬虫实现自动化数据采集，结合pandas进行多维度数据分析，并使用ECharts在前端进行丰富的可视化展示。')
add_para('本课程设计以豆瓣电影数据分析系统为题，综合运用数据库原理与应用课程所学知识，完整地经历数据库的需求分析、概念设计、逻辑设计、物理设计、实施及运行维护等全过程，深入掌握数据库应用系统的开发流程与实践方法。')

# ==================== 一、课程设计目的和要求 ====================
doc.add_heading('一、课程设计目的和要求', level=1)
add_para('课程设计是为了增强学生对所学课程的理解，学会综合地、灵活地运用所学课程知识的一个重要的实践环节。')
add_para('本课程设计应用Python程序设计语言进行数据爬取与清洗，用MySQL进行后台数据库的建立和维护，结合Navicat Premium 16进行数据库的可视化管理，编写出一个完整的电影数据分析与可视化信息系统。')
add_para('通过本课程设计可以达成如下目标：')
add_para('1. 能够自觉运用数据库原理的理论知识指导软件设计；')
add_para('2. 学会数据库的设计，并能对设计结果的优劣进行正确的评价；')
add_para('3. 学会如何组织和编写信息系统软件设计文档和软件系统的操作说明；')
add_para('4. 具有一定的独立分析问题、解决问题的能力；')
add_para('5. 掌握MySQL、Python及相关工具在数据采集、分析与可视化系统开发过程中的应用。')

# ==================== 二、课程设计过程 ====================
doc.add_heading('二、课程设计过程', level=1)
add_para('数据库设计的基本步骤如下：')
add_para('1. 需求分析')
add_para('2. 概念结构设计')
add_para('3. 数据库逻辑设计')
add_para('4. 数据库物理设计')
add_para('5. 数据库实施')
add_para('6. 数据库运行和维护')

# ---- 1. 需求分析阶段 ----
doc.add_heading('1. 需求分析阶段', level=2)

doc.add_heading('1.1 应用背景', level=3)
add_para('豆瓣电影是中国最大的电影评分与评论社区，收录了数百万部电影的基本信息，包括导演、演员、类型、评分、评论等。Top250榜单更是影迷心中的经典片单，涵盖了跨时代、跨国界的优秀电影作品。')
add_para('本系统以MySQL作为核心数据库，通过Python爬虫采集豆瓣电影Top250及各类标签页的电影数据，建立结构化的数据存储体系。在此基础上，利用SQL聚合查询和pandas数据分析库进行多维度统计分析，并通过Flask框架提供RESTful API，前端使用ECharts图表库实现丰富的可视化展示。')

doc.add_heading('1.2 需求分析目标', level=3)
add_para('1. 明确用户需求，支持电影信息的存储、查询与统计分析；')
add_para('2. 提高数据处理速度，利用数据库索引与优化技术加速查询；')
add_para('3. 实现多维度的数据可视化分析，包括类型分布、评分分布、年代趋势、国家分布等；')
add_para('4. 提供用户交互功能，支持电影浏览、收藏和评分；')
add_para('5. 保障系统的可扩展性，支持新增电影和用户数据。')

doc.add_heading('1.3 系统设计概要', level=3)
add_para('本系统主要包含四个核心模块：')
add_para('（1）数据采集模块：通过Python爬虫从豆瓣电影网页采集电影基本信息（标题、评分、年份、海报等），并通过豆瓣移动端API补充详细信息（演员、国家、语言、时长、剧情简介）。')
add_para('（2）数据存储模块（MySQL）：设计规范化的数据库表结构，包括电影表、导演表、演员表、类型表、国家表、语言表等实体表及其关联表，以及用户表、收藏表、评分表，通过建立合理的索引提升查询性能。')
add_para('（3）数据分析模块：利用pandas和SQL对电影类型分布、评分区间分布、年代趋势、国家分布、导演-演员关系网络等进行多维度统计分析。')
add_para('（4）前端展示模块：通过Flask提供RESTful API，前端使用ECharts实现9种可视化图表（饼图、柱状图、折线图、横向柱状图、词云、热力图、关系图、散点图、排行榜），支持筛选联动。')

doc.add_heading('1.4 软件处理对象', level=3)
add_para('系统要处理的对象主要包括以下方面（详细数据见数据字典）：')
add_para('1. 电影信息：包括电影ID、中文片名、原名、上映年份、豆瓣评分、评分人数、时长、上映日期、剧情简介、海报URL、豆瓣链接等，是系统核心数据对象。')
add_para('2. 导演信息：包括导演ID、姓名、豆瓣ID。一部电影可以有多个导演，一个导演可以执导多部电影。')
add_para('3. 演员信息：包括演员ID、姓名、豆瓣ID。一部电影可以有多个演员，一个演员可以出演多部电影。')
add_para('4. 类型信息：包括类型ID和类型名称（如剧情、喜剧、科幻等），与电影形成多对多关系。')
add_para('5. 国家/语言信息：记录电影的出品国家和地区、语言，与电影形成多对多关系。')
add_para('6. 用户信息：包括用户ID、用户名、密码哈希、邮箱、注册时间，用于支持收藏和评分功能。')
add_para('7. 用户收藏信息：记录用户收藏的电影，用户与电影之间的多对多联系。')
add_para('8. 用户评分信息：记录用户对电影的评分（1-5分），用户与电影之间的多对多联系。')

doc.add_heading('1.5 系统可行性分析', level=3)

doc.add_heading('1.5.1 技术可行性', level=3)
add_para('本系统采用Python + MySQL + Flask + ECharts技术栈。Python拥有成熟的requests、BeautifulSoup爬虫库和pandas数据处理库；MySQL作为广泛应用的关系型数据库，技术成熟可靠；Flask是轻量级Web框架，适合中小型项目；ECharts是百度开源的可视化库，支持丰富的图表类型。团队已系统学习数据库原理与Python编程，具备完成本系统开发的技术能力。')

doc.add_heading('1.5.2 操作可行性', level=3)
add_para('本系统提供直观的Web可视化界面和简洁的API接口，用户通过浏览器即可访问所有功能。系统采用仪表盘风格设计，图表支持交互操作（悬停提示、缩放、筛选），操作简便直观。')

doc.add_heading('结论', level=3)
add_para('经分析，本系统从技术和操作两个方面均满足可行性要求，开发本系统是完全可行的。')

doc.add_heading('1.6 系统的设计目标及其意义', level=3)
add_para('本豆瓣电影数据分析系统旨在实现以下核心目标：')
add_para('（1）可扩展性：系统支持新增电影、用户数据，数据库设计采用规范化设计，具备良好的扩展能力；')
add_para('（2）可分析性：能够对电影类型分布、评分分布、年代趋势、国家分布、导演-演员关系等进行多维度统计分析；')
add_para('（3）可视化展示：提供9种ECharts图表，支持筛选联动，直观展示数据分析结果；')
add_para('（4）用户交互：支持用户注册登录、电影浏览、收藏和评分功能；')
add_para('（5）学术价值：系统具备完整的E-R图、流程图、SQL示例和可视化结果，具有较好的学术展示价值。')

doc.add_heading('1.7 系统的业务流程及具体的功能', level=3)
add_para('通过对用户需求和系统设计思想的分析，本系统大致可以分为以下几大模块：数据采集模块、数据存储模块、数据分析模块、前端展示模块。')

add_image_placeholder('系统业务流程图')

add_image_placeholder('系统功能模块图')

add_para('主要模块的功能：')
add_para('①数据采集模块：通过Python爬虫采集豆瓣电影Top250和标签页电影列表，再通过豆瓣移动端API获取详细信息（演员、国家、语言、时长、简介），海报图片批量下载到本地存储。')
add_para('②数据存储模块：设计并维护movies、directors、actors、genres、countries、languages等实体表及movie_directors、movie_actors、movie_genres等关联表，以及users、user_favorites、user_ratings用户相关表，建立外键约束和索引。')
add_para('③数据分析模块：利用SQL聚合查询和pandas对电影类型分布、评分区间分布、年代趋势、国家分布、导演-演员网络、关键词词频等进行分析。')
add_para('④前端展示模块：通过Flask提供RESTful API，前端使用ECharts实现饼图、柱状图、折线图、横向柱状图、词云、热力图、关系图、散点图、排行榜共9种图表，支持类型/年份/国家/评分筛选联动。')

doc.add_heading('1.8 数据流程', level=3)
add_para('系统的数据流程如下：')
add_para('（1）数据采集：Python爬虫从豆瓣电影网页抓取电影列表和基本信息，通过移动端API补充详细数据；')
add_para('（2）数据清洗与入库：对采集的数据进行去重、格式转换、空值处理，通过参数化SQL写入MySQL数据库；')
add_para('（3）数据分析：后端通过SQL聚合查询和pandas对数据库中的数据进行多维度统计分析；')
add_para('（4）API服务：Flask将分析结果封装为JSON格式的RESTful API接口；')
add_para('（5）前端展示：浏览器通过API获取数据，使用ECharts渲染为交互式可视化图表。')

add_image_placeholder('系统数据流程图')

# ---- 2. 数据字典 ----
doc.add_heading('2. 系统的数据字典', level=2)
add_para('数据字典是对系统中数据的详细说明，包含各数据表的属性名、类型、约束和说明信息。')

add_para('电影数据字典（movies）：', bold=True)
add_table(
    ['属性名', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PRIMARY KEY AUTO_INCREMENT', '主键'],
        ['douban_id', 'VARCHAR(20)', 'UNIQUE NOT NULL', '豆瓣电影ID'],
        ['title', 'VARCHAR(255)', 'NOT NULL', '中文片名'],
        ['original_title', 'VARCHAR(255)', '', '原名'],
        ['year', 'INT', 'INDEX', '上映年份'],
        ['rating', 'DECIMAL(3,1)', 'INDEX', '豆瓣评分'],
        ['rating_count', 'INT', '', '评分人数'],
        ['duration_minutes', 'INT', '', '时长（分钟）'],
        ['release_date', 'DATE', '', '上映日期'],
        ['summary', 'TEXT', '', '剧情简介'],
        ['poster_url', 'VARCHAR(500)', '', '海报图片URL'],
        ['douban_url', 'VARCHAR(500)', '', '豆瓣链接'],
        ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '入库时间'],
        ['updated_at', 'TIMESTAMP', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
    ]
)

doc.add_paragraph()
add_para('导演数据字典（directors）：', bold=True)
add_table(
    ['属性名', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PRIMARY KEY AUTO_INCREMENT', '主键'],
        ['name', 'VARCHAR(100)', 'NOT NULL', '导演姓名'],
        ['douban_id', 'VARCHAR(20)', 'UNIQUE', '豆瓣ID'],
    ]
)

doc.add_paragraph()
add_para('演员数据字典（actors）：', bold=True)
add_para('结构同 directors 表。')

doc.add_paragraph()
add_para('类型数据字典（genres）：', bold=True)
add_table(
    ['属性名', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PRIMARY KEY AUTO_INCREMENT', '主键'],
        ['name', 'VARCHAR(50)', 'UNIQUE NOT NULL', '类型名称（如剧情、喜剧、科幻）'],
    ]
)

doc.add_paragraph()
add_para('国家/地区数据字典（countries）与语言数据字典（languages）：', bold=True)
add_para('结构同 genres 表。')

doc.add_paragraph()
add_para('电影-导演关联表（movie_directors）：', bold=True)
add_table(
    ['属性名', '类型', '约束', '说明'],
    [
        ['movie_id', 'INT', 'FOREIGN KEY → movies(id)', '电影ID'],
        ['director_id', 'INT', 'FOREIGN KEY → directors(id)', '导演ID'],
        ['', '', 'PRIMARY KEY (movie_id, director_id)', '联合主键'],
    ]
)

doc.add_paragraph()
add_para('电影-演员关联表（movie_actors）：', bold=True)
add_table(
    ['属性名', '类型', '约束', '说明'],
    [
        ['movie_id', 'INT', 'FOREIGN KEY → movies(id)', '电影ID'],
        ['actor_id', 'INT', 'FOREIGN KEY → actors(id)', '演员ID'],
        ['order_no', 'INT', '', '主演顺序'],
        ['', '', 'PRIMARY KEY (movie_id, actor_id)', '联合主键'],
    ]
)

doc.add_paragraph()
add_para('电影-类型关联表（movie_genres）、电影-国家关联表（movie_countries）、电影-语言关联表（movie_languages）：', bold=True)
add_para('结构类似 movie_directors，均为 (movie_id, foreign_id) 联合主键，外键级联删除。')

doc.add_paragraph()
add_para('用户数据字典（users）：', bold=True)
add_table(
    ['属性名', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PRIMARY KEY AUTO_INCREMENT', '主键'],
        ['username', 'VARCHAR(50)', 'UNIQUE NOT NULL', '用户名'],
        ['password_hash', 'VARCHAR(255)', 'NOT NULL', '密码哈希'],
        ['email', 'VARCHAR(100)', '', '邮箱（可选）'],
        ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '注册时间'],
    ]
)

doc.add_paragraph()
add_para('用户收藏数据字典（user_favorites）：', bold=True)
add_table(
    ['属性名', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PRIMARY KEY AUTO_INCREMENT', '主键'],
        ['user_id', 'INT', 'FOREIGN KEY → users(id)', '用户ID'],
        ['movie_id', 'INT', 'FOREIGN KEY → movies(id)', '电影ID'],
        ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '收藏时间'],
        ['', '', 'UNIQUE (user_id, movie_id)', '联合唯一索引'],
    ]
)

doc.add_paragraph()
add_para('用户评分数据字典（user_ratings）：', bold=True)
add_table(
    ['属性名', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PRIMARY KEY AUTO_INCREMENT', '主键'],
        ['user_id', 'INT', 'FOREIGN KEY → users(id)', '用户ID'],
        ['movie_id', 'INT', 'FOREIGN KEY → movies(id)', '电影ID'],
        ['rating', 'DECIMAL(2,1)', 'NOT NULL', '用户评分（1.0-5.0）'],
        ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '评分时间'],
        ['updated_at', 'TIMESTAMP', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
        ['', '', 'UNIQUE (user_id, movie_id)', '联合唯一索引'],
    ]
)

# ---- 3. 概念结构设计 ----
doc.add_heading('3. 概念结构设计阶段', level=2)
add_para('本系统的总E-R图如下图所示：')
add_image_placeholder('系统总E-R图')

add_para('以下是各实体的分E-R图描述：')
add_para('电影实体（movies）：主要属性包括id、douban_id、title、original_title、year、rating、rating_count、duration_minutes、release_date、summary、poster_url、douban_url。电影是系统核心实体，与导演、演员、类型、国家、语言形成多对多关系。')
add_para('导演实体（directors）：主要属性包括id、name、douban_id。导演与电影之间是多对多关系（一个导演可执导多部电影，一部电影可有多个导演）。')
add_para('演员实体（actors）：主要属性包括id、name、douban_id。演员与电影之间是多对多关系。')
add_para('类型实体（genres）：主要属性包括id、name。类型与电影之间是多对多关系。')
add_para('国家实体（countries）与语言实体（languages）：结构同类型实体，与电影之间均为多对多关系。')
add_para('用户实体（users）：主要属性包括id、username、password_hash、email。用户与电影之间通过收藏和评分形成多对多关系。')

add_image_placeholder('电影-导演-演员分E-R图')

# ---- 4. 逻辑结构设计 ----
doc.add_heading('4. 逻辑结构设计阶段', level=2)
add_para('将系统的E-R图转换成关系模型如下：')
add_para('movies(id, douban_id, title, original_title, year, rating, rating_count, duration_minutes, release_date, summary, poster_url, douban_url, created_at, updated_at)')
add_para('directors(id, name, douban_id)')
add_para('actors(id, name, douban_id)')
add_para('genres(id, name)')
add_para('countries(id, name)')
add_para('languages(id, name)')
add_para('movie_directors(movie_id, director_id)')
add_para('movie_actors(movie_id, actor_id, order_no)')
add_para('movie_genres(movie_id, genre_id)')
add_para('movie_countries(movie_id, country_id)')
add_para('movie_languages(movie_id, language_id)')
add_para('users(id, username, password_hash, email, created_at)')
add_para('user_favorites(id, user_id, movie_id, created_at)')
add_para('user_ratings(id, user_id, movie_id, rating, created_at, updated_at)')
add_para('其中，movie_directors、movie_actors、movie_genres、movie_countries、movie_languages中的movie_id为外码，参照movies表的id；对应的director_id、actor_id、genre_id、country_id、language_id分别参照各自的实体表。所有关联表使用ON DELETE CASCADE级联删除策略。')
add_para('user_favorites和user_ratings中的user_id为外码参照users表的id，movie_id为外码参照movies表的id。')
add_para('本系统的关系模式均满足第三范式（3NF），每个非主属性完全函数依赖于主码，且不存在非主属性对主码的传递依赖，数据冗余度低，更新异常少。')

# ---- 5. 物理结构设计 ----
doc.add_heading('5. 物理结构设计阶段', level=2)
add_para('数据库的物理设计通常分为两步：（1）确定数据库的物理结构，主要指存取方法和存储结构；（2）对物理结构进行评价，评价的重点是时间和空间的效率。')

doc.add_heading('5.1 关系模式存取方法的选择', level=3)
add_para('本系统在经常查询和连接的列上建立索引，以提高查询性能。具体策略如下：')
add_para('（1）movies表建立douban_id的UNIQUE索引，加速按豆瓣ID的查询和去重判断；')
add_para('（2）movies表建立year和rating的单列索引，加速按年份筛选和评分排序查询；')
add_para('（3）关联表使用联合主键(movie_id, foreign_id)，数据库自动创建主键索引，加速连接查询；')
add_para('（4）user_favorites和user_ratings表建立(user_id, movie_id)的联合唯一索引，加速用户收藏/评分状态查询。')

doc.add_heading('5.2 确定数据库的存储结构', level=3)
add_para('本系统采用MySQL 8.0数据库，使用InnoDB存储引擎，支持事务处理和外键约束。数据库字符集设置为utf8mb4，排序规则为utf8mb4_unicode_ci，确保中文数据的正确存储和排序。数据文件和日志文件存放于MySQL默认数据目录。')

# ---- 6. 数据库实施 ----
doc.add_heading('6. 数据库实施', level=2)
add_para('本系统前端使用Python进行数据采集与处理，后台数据库采用MySQL。使用Python的mysql-connector-python库进行数据库连接与操作（不使用ORM，与课程教学保持一致）。')
add_para('实施环境采用Python 3.10+、MySQL 8.0和Navicat Premium 16。数据入库前先进行去重、空值检查和字段类型转换，确保电影、导演、演员、类型等核心数据能够按照主外键关系稳定写入数据库。')

doc.add_heading('6.1 Create Database 创建数据库', level=3)
add_code('-- 创建数据库：豆瓣电影数据分析系统')
add_code('CREATE DATABASE douban_movies')
add_code('  CHARACTER SET utf8mb4')
add_code('  COLLATE utf8mb4_unicode_ci;')
add_code('USE douban_movies;')

doc.add_heading('6.2 Create Table 创建表', level=3)
add_para('创建电影表 movies：')
add_code('''CREATE TABLE movies (
    id              INT          NOT NULL AUTO_INCREMENT PRIMARY KEY,
    douban_id       VARCHAR(20)  NOT NULL,
    title           VARCHAR(255) NOT NULL,
    original_title  VARCHAR(255) DEFAULT NULL,
    year            INT          DEFAULT NULL,
    rating          DECIMAL(3,1) DEFAULT NULL,
    rating_count    INT          DEFAULT NULL,
    duration_minutes INT         DEFAULT NULL,
    release_date    DATE         DEFAULT NULL,
    summary         TEXT,
    poster_url      VARCHAR(500) DEFAULT NULL,
    douban_url      VARCHAR(500) DEFAULT NULL,
    created_at      TIMESTAMP    DEFAULT CURRENT_TIMESTAMP,
    updated_at      TIMESTAMP    DEFAULT NULL ON UPDATE CURRENT_TIMESTAMP,
    UNIQUE KEY uk_douban_id (douban_id),
    INDEX idx_year (year),
    INDEX idx_rating (rating)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

add_para('创建导演表 directors：')
add_code('''CREATE TABLE directors (
    id        INT          NOT NULL AUTO_INCREMENT PRIMARY KEY,
    name      VARCHAR(100) NOT NULL,
    douban_id VARCHAR(20)  DEFAULT NULL,
    UNIQUE KEY uk_douban_id (douban_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

add_para('创建演员表 actors：')
add_code('（结构同 directors 表）')

add_para('创建类型表 genres：')
add_code('''CREATE TABLE genres (
    id   INT         NOT NULL AUTO_INCREMENT PRIMARY KEY,
    name VARCHAR(50) NOT NULL,
    UNIQUE KEY uk_name (name)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

add_para('创建国家表 countries 和语言表 languages：')
add_code('（结构同 genres 表）')

add_para('创建电影-导演关联表 movie_directors：')
add_code('''CREATE TABLE movie_directors (
    movie_id    INT NOT NULL,
    director_id INT NOT NULL,
    PRIMARY KEY (movie_id, director_id),
    FOREIGN KEY (movie_id)    REFERENCES movies(id)    ON DELETE CASCADE,
    FOREIGN KEY (director_id) REFERENCES directors(id)  ON DELETE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

add_para('创建电影-演员关联表 movie_actors：')
add_code('''CREATE TABLE movie_actors (
    movie_id  INT NOT NULL,
    actor_id  INT NOT NULL,
    order_no  INT DEFAULT NULL,
    PRIMARY KEY (movie_id, actor_id),
    FOREIGN KEY (movie_id) REFERENCES movies(id) ON DELETE CASCADE,
    FOREIGN KEY (actor_id) REFERENCES actors(id) ON DELETE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

add_para('创建电影-类型关联表 movie_genres：')
add_code('''CREATE TABLE movie_genres (
    movie_id INT NOT NULL,
    genre_id INT NOT NULL,
    PRIMARY KEY (movie_id, genre_id),
    FOREIGN KEY (movie_id) REFERENCES movies(id) ON DELETE CASCADE,
    FOREIGN KEY (genre_id) REFERENCES genres(id) ON DELETE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

add_para('创建电影-国家关联表 movie_countries 和电影-语言关联表 movie_languages：')
add_code('（结构同 movie_genres，外键分别引用 countries 和 languages 表）')

add_para('创建用户表 users：')
add_code('''CREATE TABLE users (
    id            INT          NOT NULL AUTO_INCREMENT PRIMARY KEY,
    username      VARCHAR(50)  NOT NULL,
    password_hash VARCHAR(255) NOT NULL,
    email         VARCHAR(100) DEFAULT NULL,
    created_at    TIMESTAMP    DEFAULT CURRENT_TIMESTAMP,
    UNIQUE KEY uk_username (username)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

add_para('创建用户收藏表 user_favorites：')
add_code('''CREATE TABLE user_favorites (
    id         INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
    user_id    INT NOT NULL,
    movie_id   INT NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    UNIQUE KEY uk_user_movie (user_id, movie_id),
    FOREIGN KEY (user_id)  REFERENCES users(id)  ON DELETE CASCADE,
    FOREIGN KEY (movie_id) REFERENCES movies(id) ON DELETE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

add_para('创建用户评分表 user_ratings：')
add_code('''CREATE TABLE user_ratings (
    id         INT          NOT NULL AUTO_INCREMENT PRIMARY KEY,
    user_id    INT          NOT NULL,
    movie_id   INT          NOT NULL,
    rating     DECIMAL(2,1) NOT NULL,
    created_at TIMESTAMP    DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP    DEFAULT NULL ON UPDATE CURRENT_TIMESTAMP,
    UNIQUE KEY uk_user_movie (user_id, movie_id),
    FOREIGN KEY (user_id)  REFERENCES users(id)  ON DELETE CASCADE,
    FOREIGN KEY (movie_id) REFERENCES movies(id) ON DELETE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;''')

doc.add_heading('6.3 关键代码展示', level=3)

add_para('（1）参数化SQL查询（防SQL注入）', bold=True)
add_para('系统所有数据库操作均使用参数化查询，杜绝SQL注入风险：')
add_code('''def get_movies(page, size, genre=None, year=None, min_rating=None, country=None):
    """获取电影列表（支持筛选）"""
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    conditions, params = _build_filter_conditions(genre, year, min_rating, country)
    where = "WHERE " + " AND ".join(conditions) if conditions else ""
    sql = f"SELECT SQL_CALC_FOUND_ROWS m.* FROM movies m {where}
           ORDER BY m.rating DESC LIMIT %s OFFSET %s"
    params.extend([size, (page - 1) * size])
    cursor.execute(sql, params)
    movies = cursor.fetchall()''')

add_para('（2）爬虫反爬策略', bold=True)
add_para('系统采用多种反爬措施，确保数据采集的稳定性：')
add_code('''# User-Agent 池轮换
USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 ...",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 ...",
    # ... 共10+个UA
]

# 请求间隔随机化 + 指数退避重试
time.sleep(random.uniform(CRAWL_MIN_INTERVAL, CRAWL_MAX_INTERVAL))
for retry in range(CRAWL_MAX_RETRIES):
    response = requests.get(url, headers=headers, timeout=10)
    if response.status_code == 200:
        break
    time.sleep(CRAWL_RETRY_DELAYS[retry])  # 2s, 4s, 8s''')

add_para('（3）海报防盗链处理', bold=True)
add_para('豆瓣CDN要求请求携带正确的Referer头，否则返回HTTP 418。解决方案是将海报批量下载到本地存储：')
add_code('''# Flask海报代理接口（备选方案）
@app.route('/api/proxy/poster')
def proxy_poster():
    url = request.args.get('url', '')
    headers = {
        'Referer': 'https://movie.douban.com/',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) ...'
    }
    resp = requests.get(url, headers=headers, timeout=5)
    return Response(resp.content, content_type='image/jpeg')''')

# ---- 7. 数据库运行和维护 ----
doc.add_heading('7. 数据库运行和维护', level=2)

doc.add_heading('7.1 解决问题的方法', level=3)
add_para('在系统开发过程中，遇到并解决了以下主要问题：')
add_para('（1）豆瓣详情页反爬：豆瓣电影详情页使用SHA-512工作量证明（PoW），Python爬虫无法直接解析。解决方案是改用豆瓣移动端API（m.douban.com/rexxar/api/v2/movie/{id}）获取JSON格式的详细数据。')
add_para('（2）海报图片防盗链：豆瓣CDN检查Referer头，非豆瓣域名返回HTTP 418。解决方案是批量下载海报到本地static/posters/目录，并提供Flask代理接口作为备选。')
add_para('（3）前端性能优化：初始版本每张电影卡片单独请求收藏和评分状态（24个请求），优化为批量获取后传入（3个请求），并添加Tab切换缓存和海报懒加载。')
add_para('（4）ECharts CDN不可用：jsdelivr等国际CDN在国内访问超时，将ECharts和词云扩展下载到本地static/js/lib/目录。')

doc.add_heading('7.2 系统维护', level=3)
add_para('系统日常维护工作包括：')
add_para('（1）数据更新：可通过运行爬虫脚本增量采集新电影数据，已存在的电影通过douban_id自动跳过；')
add_para('（2）索引维护：MySQL InnoDB引擎自动维护B+树索引，无需手动重建；')
add_para('（3）备份策略：通过mysqldump定期备份数据库，确保数据安全；')
add_para('（4）海报管理：poster_url字段存储本地路径（/static/posters/xxx.jpg），新增电影可通过scripts/download_posters.py脚本补充下载。')

doc.add_heading('7.3 数据库性能的评价', level=3)
add_para('本系统当前数据库包含250部电影、280位导演、4953位演员、15种类型、32个国家/地区的完整数据。主要查询性能表现如下：')
add_para('（1）电影列表查询（带筛选条件）：通过year、rating索引，响应时间在毫秒级；')
add_para('（2）多表连接查询（电影+导演+演员+类型）：通过联合主键索引，响应时间 < 50ms；')
add_para('（3）统计分析查询（类型分布、评分分布等）：利用SQL聚合函数和索引，单次查询 < 100ms；')
add_para('（4）前端加载性能：优化后首次加载仅需3个并行API请求（原需24个串行请求），Tab切换命中缓存后零网络请求。')

# ==================== 三、课程设计心得 ====================
doc.add_heading('三、课程设计心得', level=1)
add_para('通过本次课程设计，我完整地经历了一个数据库应用系统从需求分析到实施运行的全过程，对数据库原理课程所学的理论知识有了更加深刻的理解和实践体会。')
add_para('在需求分析阶段，我学会了如何从实际应用场景出发，梳理系统的数据需求和功能需求，绘制数据流程图和业务流程图。在概念设计阶段，通过绘制E-R图，我更加深入地理解了实体、属性、联系的概念，以及如何将现实世界的信息结构抽象为概念模型。')
add_para('在逻辑设计和物理设计阶段，我掌握了将E-R模型转换为关系模型的方法，理解了第三范式的意义和作用，并学会了通过建立索引来优化查询性能。特别是在多对多关系的处理上，关联表+外键+级联删除的设计模式让我对数据库的完整性约束有了直观的认识。')
add_para('在实施阶段，我深刻体会到参数化SQL的重要性——它不仅防止了SQL注入攻击，也使代码更加清晰可维护。在数据采集过程中，面对豆瓣网站的反爬机制（SHA-512 PoW工作量证明、CDN防盗链、请求频率限制），我学会了通过分析HTTP请求、使用移动端API、伪造请求头等多种技术手段来解决实际问题。')
add_para('此外，本次课程设计还让我接触到了Web后端开发（Flask框架）、前端可视化（ECharts图表库）、数据处理（pandas）等数据库之外的技术，拓宽了我的技术视野。整个开发过程锻炼了我独立分析问题、解决问题的能力，也让我认识到理论学习与实践应用之间的差距——很多在课本上看似简单的概念，在实际开发中可能会遇到各种意想不到的问题，只有通过动手实践才能真正掌握。')

# ==================== 保存 ====================
doc.save(OUTPUT)
print(f'报告已生成: {OUTPUT}')
